import os, filecmp
from difflibparser.difflibparser import *
from ui.mainwindow_ui import MainWindowUI
from docx import Document
import zipfile
import time
import pdfplumber #import fitz , fitz can be used to extract text from pdfs



from tkinter import *
from tkinter.filedialog import askopenfilename, askdirectory
from tkinter.simpledialog import askstring

def unzip_all(directory):
        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith('.zip'):
                    file_path = os.path.join(root, file)
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        zip_ref.extractall(root)


class MainWindow:
    def start(self, leftpath = None, rightpath = None):
        self.main_window = Tk()
        self.main_window.title('difference detection')
        self.__main_window_ui = MainWindowUI(self.main_window)

        self.leftFile = ''
        self.rightFile = ''

        self.__main_window_ui.center_window()
        self.__main_window_ui.create_file_path_labels()
        self.__main_window_ui.create_text_areas()
        self.__main_window_ui.create_search_text_entry(self.__findNext)
        self.__main_window_ui.create_line_numbers()
        self.__main_window_ui.create_scroll_bars()
        self.__main_window_ui.create_file_treeview()
        path_to_my_project = os.getcwd()
        self.__main_window_ui.add_menu('File', [
            {'name': 'Compare Files', 'command': self.__browse_files},
            {'name': 'Compare Directories', 'command': self.__browse_directories},
            {'separator'},
            {'name': 'Exit', 'command': self.__exit, 'accelerator': 'Alt+F4'}
            ])
        self.__main_window_ui.add_menu('Edit', [
            {'name': 'Find', 'command': self.__startFindText, 'accelerator': 'Ctrl+F'},
            {'separator'},
            {'name': 'Cut', 'command': self.__cut, 'accelerator': 'Ctrl+X'},
            {'name': 'Copy', 'command': self.__copy, 'accelerator': 'Ctrl+C'},
            {'name': 'Paste', 'command': self.__paste, 'accelerator': 'Ctrl+V'},
            {'separator'},
            {'name': 'Go To Line', 'command': self.__goToLine, 'accelerator': 'Ctrl+G'}
            ])
        self.__main_window_ui.fileTreeView.bind('<<TreeviewSelect>>', lambda *x:self.treeViewItemSelected())

        if (leftpath and os.path.isdir(leftpath)) or (rightpath and os.path.isdir(rightpath)):
            self.__load_directories(leftpath, rightpath)
        else:
            self.leftFile = leftpath if leftpath else ''
            self.rightFile = rightpath if rightpath else ''
            self.filesChanged()

        self.__bind_key_shortcuts()

        self.main_window.mainloop()

    def __bind_key_shortcuts(self):
        self.main_window.bind('<Control-f>', lambda *x: self.__startFindText())
        self.main_window.bind('<Control-g>', lambda *x: self.__goToLine())
        self.main_window.bind('<Escape>', lambda *x: self.__endFindText())
        self.main_window.bind('<F3>', self.__main_window_ui.searchTextDialog.nextResult)

    def __browse_files(self):
        self.__load_file('left')
        self.__load_file('right')
        self.filesChanged()
        self.__main_window_ui.fileTreeView.grid_remove()
        self.__main_window_ui.fileTreeYScrollbar.grid_remove()
        self.__main_window_ui.fileTreeXScrollbar.grid_remove()

    # Load directories into the treeview
    def __browse_directories(self):
        leftDir = self.__load_directory('left')
        rightDir = self.__load_directory('right')
        self.__load_directories(leftDir, rightDir)

    def __load_directories(self, leftDir, rightDir):
        if leftDir and rightDir:
            self.__main_window_ui.fileTreeView.grid()
            self.__main_window_ui.fileTreeYScrollbar.grid()
            self.__main_window_ui.fileTreeXScrollbar.grid()
            self.__main_window_ui.fileTreeView.delete(*self.__main_window_ui.fileTreeView.get_children())
            self.__browse_process_directory('', leftDir, rightDir)
    
        
    def __convert_pdf_to_txt(self, pdf_file, output_file):
        with pdfplumber.open(pdf_file) as pdf:
            with open(output_file, 'w') as f:
                for page in pdf.pages:
                    # Extract text 
                    text = page.extract_text()
                    if text:
                        f.write(text.strip() + '\n')
                    else:
                        f.write("Empty PDF page")

                    # Extract table 
                    tables = page.extract_tables()
                    for table in tables:
                        f.write('\n' + '-' * 100 + '\n')  

                        # Determine the number of columns
                        num_columns = max(len(row) for row in table)

                        # Process each row
                        for row in table:
                            row_data = [''] * num_columns
                            for i, cell in enumerate(row):
                                
                                cell_content = (cell or '').replace('\n', ' ').strip()

                                if i < num_columns:
                                    row_data[i] = cell_content

                            
                            formatted_row = ' | '.join(cell.ljust(20) for cell in row_data)
                            f.write(formatted_row + '\n')

                        f.write('-' * 100 + '\n')  

                    f.write('\n')


    
    """
    
    def __convert_docx_to_txt(self, docx_file, output_file):
        doc = Document(docx_file)

        with open(output_file, 'w') as f:
            if doc.paragraphs:
                for para in doc.paragraphs:
                    if para.text.strip():
                        f.write(para.text + '\n')
            else:
                f.write("Empty DOCX file\n")

            if doc.tables:
                for table in doc.tables:
                    num_columns = max(len(row.cells) for row in table.rows)
                    table_width = num_columns * 17  
                    separator_line = '-' * table_width

                    f.write(separator_line + '\n')  
                    for row_idx, row in enumerate(table.rows):
                        row_data = [''] * num_columns  
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.replace('\n', ' ')
                            top = cell._tc.top
                            bottom = cell._tc.bottom
                            left = cell._tc.left
                            right = cell._tc.right

                            if top == row_idx and left == col_idx:
                                span = (right - left) * 15
                                cell_text = f'{cell_text:<{span}}'
                                row_data[left] = cell_text 
                            elif top == row_idx:
                                row_data[left] = ''  

                        formatted_row = ' | '.join(row_data)
                        f.write(formatted_row + '\n')
                        if row_idx == len(table.rows) - 1 or table.rows[row_idx + 1].cells[col_idx]._tc.top != row_idx:
                            f.write(separator_line + '\n')  

                    f.write('\n')
                    
   """ 
    
    def __convert_docx_to_txt(self, docx_file, output_file):
        doc = Document(docx_file)

        def write_paragraph(paragraph, file):
            if paragraph.text.strip():
                file.write(paragraph.text + '\n')

        def write_table(table, file):
            num_columns = max(len(row.cells) for row in table.rows)
            cell_width = 20  
            separator_line = '-' * (num_columns * cell_width + (num_columns - 1) * 3)


            # Add a line before each table
            file.write(separator_line + '\n') 

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.replace('\n', ' ').strip()
                    row_data.append(cell_text.ljust(cell_width))

                file.write(' | '.join(row_data) + '\n')

            # Add a line after each table
            file.write(separator_line + '\n')  

        with open(output_file, 'w') as f:
            paragraph_idx = 0
            table_idx = 0

            for element in doc.element.body:

                # extract text
                if element.tag.endswith('p'):  
                    paragraph = doc.paragraphs[paragraph_idx]
                    write_paragraph(paragraph, f)
                    paragraph_idx += 1


                # extract table
                elif element.tag.endswith('tbl'):  
                    table = doc.tables[table_idx]
                    write_table(table, f)
                    table_idx += 1

            if paragraph_idx == 0 and table_idx == 0:
                f.write("Empty DOCX file")


    def __browse_process_directory(self, parent, leftPath, rightPath):
        unzip_all(leftPath)
        unzip_all(rightPath)
       
        if parent == '':
            leftPath = leftPath.rstrip('/')
            rightPath = rightPath.rstrip('/')
            leftDirName = os.path.basename(leftPath)
            rightDirName = os.path.basename(rightPath)
            self.__main_window_ui.fileTreeView.heading('#0', text=leftDirName + ' / ' + rightDirName, anchor=W)
       
        conversions_complete = False
        
        
        while not conversions_complete:
            leftListing = os.listdir(leftPath)
            rightListing = os.listdir(rightPath)
            mergedListing = list(set(leftListing) | set(rightListing))
            painted = FALSE
            total_conversions_docx = sum(1 for l in mergedListing if l.endswith('.docx'))
            total_conversions_pdf = sum(1 for l in mergedListing if l.endswith('.pdf'))
            num_conversions_docx = 0
            num_conversions_pdf = 0
    
            
            for l in mergedListing:
                newLeftPath = os.path.join(leftPath, l)
                newRightPath = os.path.join(rightPath, l)
                
                if l.endswith('.docx'):
                    num_conversions_docx += 1
                    if l in leftListing:
                        txt_file = os.path.splitext(l)[0] + '_docx_converted.txt'
                        output_path = os.path.join(leftPath, txt_file)
                        self.__convert_docx_to_txt(newLeftPath, output_path)
                    if l in rightListing:
                        txt_file = os.path.splitext(l)[0] + '_docx_converted.txt'
                        output_path = os.path.join(rightPath, txt_file)
                        self.__convert_docx_to_txt(newRightPath, output_path)
                

                if l.endswith('.pdf'):
                            num_conversions_pdf += 1
                            if l in leftListing:
                                txt_file = os.path.splitext(l)[0] + '_pdf_converted.txt'
                                output_path = os.path.join(leftPath, txt_file)
                                self.__convert_pdf_to_txt(newLeftPath, output_path)
                            if l in rightListing:
                                txt_file = os.path.splitext(l)[0] + '_pdf_converted.txt'
                                output_path = os.path.join(rightPath, txt_file)
                                self.__convert_pdf_to_txt(newRightPath, output_path)


            if num_conversions_docx == total_conversions_docx and num_conversions_pdf == total_conversions_pdf:
                conversions_complete = True
            else:
                time.sleep(0.1)

        leftListing = os.listdir(leftPath)
        rightListing = os.listdir(rightPath)
        mergedListing = list(set(leftListing) | set(rightListing))
        
        for l in mergedListing:
            
            newLeftPath = leftPath + '/' + l
            newRightPath = rightPath + '/' + l
            bindValue = (newLeftPath, newRightPath)
           
            if not (l.endswith('docx') or l.endswith('zip') or l.endswith('pdf') or l.endswith('svn')) :
                # Item in left dir only
                if l in leftListing and l not in rightListing:
                    self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, value=bindValue, open=False, tags=('red','simple'))
                    painted = TRUE
                # Item in right dir only
                elif l in rightListing and l not in leftListing:
                    self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, value=bindValue, open=False, tags=('green','simple'))
                    painted = TRUE
                # Item in both dirs
                else:
                    # If one of the diffed items is a file and the other is a directory, show in yellow indicating a difference
                    if (not os.path.isdir(newLeftPath) and os.path.isdir(newRightPath)) or (os.path.isdir(newLeftPath) and not os.path.isdir(newRightPath)):
                        self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, value=bindValue, open=False, tags=('yellow','simple'))
                        painted = TRUE
                    else:
                        # If both are directories, show in white and recurse on contents
                        if os.path.isdir(newLeftPath) and os.path.isdir(newRightPath):
                            oid = self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, open=False)
                            painted = self.__browse_process_directory(oid, newLeftPath, newRightPath)
                            if painted:
                                self.__main_window_ui.fileTreeView.item(oid, tags=('purpleLight', 'simple'))
                        else:
                            # Both are files. diff the two files to either show them in white or yellow
                            if (filecmp.cmp(newLeftPath, newRightPath)):
                                oid = self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, value=bindValue, open=False, tags=('simple'))
                            else:
                                oid = self.__main_window_ui.fileTreeView.insert(parent, 'end', text=l, value=bindValue, open=False, tags=('yellow','simple'))
                                painted = TRUE
        return painted



    

    def __load_file(self, pos):
        fname = askopenfilename()
        if fname:
            if fname.endswith('.docx'):
            
                txt_file = os.path.splitext(fname)[0] + '_docx_converted.txt'
                self.__convert_docx_to_txt(fname, txt_file)
                if pos == 'left':
                    self.leftFile = txt_file
                else:
                    self.rightFile = txt_file

            elif fname.endswith('.pdf'):
            
                txt_file = os.path.splitext(fname)[0] + '_pdf_converted.txt'
                self.__convert_pdf_to_txt(fname, txt_file)
                if pos == 'left':
                    self.leftFile = txt_file
                else:
                    self.rightFile = txt_file

            else:
                # Load TXT file
                if pos == 'left':
                    self.leftFile = fname
                else:
                    self.rightFile = fname
            return fname
        else:
            return None

    def __load_directory(self, pos):
        dirName = askdirectory()
        if dirName:
            if pos == 'left':
                self.__main_window_ui.leftFileLabel.config(text=dirName)
            else:
                self.__main_window_ui.rightFileLabel.config(text=dirName)
            return dirName
        else:
            return None

    # Callback for changing a file path
    def filesChanged(self):
        self.__main_window_ui.leftLinenumbers.grid_remove()
        self.__main_window_ui.rightLinenumbers.grid_remove()

        if not self.leftFile or not self.rightFile:
            self.__main_window_ui.leftFileTextArea.config(background=self.__main_window_ui.grayColor)
            self.__main_window_ui.rightFileTextArea.config(background=self.__main_window_ui.grayColor)
            return

        if os.path.exists(self.leftFile):
            self.__main_window_ui.leftFileLabel.config(text=self.leftFile)
            self.__main_window_ui.leftFileTextArea.config(background=self.__main_window_ui.whiteColor)
            self.__main_window_ui.leftLinenumbers.grid()
        else:
            self.__main_window_ui.leftFileLabel.config(text='')

        if os.path.exists(self.rightFile):
            self.__main_window_ui.rightFileLabel.config(text=self.rightFile)
            self.__main_window_ui.rightFileTextArea.config(background=self.__main_window_ui.whiteColor)
            self.__main_window_ui.rightLinenumbers.grid()
        else:
            self.__main_window_ui.rightFileLabel.config(text='')

        self.diff_files_into_text_areas()

    def treeViewItemSelected(self):
        item_id = self.__main_window_ui.fileTreeView.focus()
        paths = self.__main_window_ui.fileTreeView.item(item_id)['values']
        if paths == None or len(paths) == 0:
            return
        self.leftFile = paths[0]
        self.rightFile = paths[1]
        self.filesChanged()

    # Insert file contents into text areas and highlight differences
    def diff_files_into_text_areas(self):
        try:
            leftFileContents = open(self.leftFile).read()
        except:
            leftFileContents = ''
        try:
            rightFileContents = open(self.rightFile).read()
        except:
            rightFileContents = ''

        diff = DifflibParser(leftFileContents.splitlines(), rightFileContents.splitlines())

        # enable text area edits so we can clear and insert into them
        self.__main_window_ui.leftFileTextArea.config(state=NORMAL)
        self.__main_window_ui.rightFileTextArea.config(state=NORMAL)
        self.__main_window_ui.leftLinenumbers.config(state=NORMAL)
        self.__main_window_ui.rightLinenumbers.config(state=NORMAL)

        # clear all text areas
        self.__main_window_ui.leftFileTextArea.delete(1.0, END)
        self.__main_window_ui.rightFileTextArea.delete(1.0, END)
        self.__main_window_ui.leftLinenumbers.delete(1.0, END)
        self.__main_window_ui.rightLinenumbers.delete(1.0, END)

        leftlineno = rightlineno = 1
        for line in diff:
            if line['code'] == DiffCode.SIMILAR:
                self.__main_window_ui.leftFileTextArea.insert('end', line['line'] + '\n')
                self.__main_window_ui.rightFileTextArea.insert('end', line['line'] + '\n')
            elif line['code'] == DiffCode.RIGHTONLY:
                self.__main_window_ui.leftFileTextArea.insert('end', '\n', 'gray')
                self.__main_window_ui.rightFileTextArea.insert('end', line['line'] + '\n', 'green')
            elif line['code'] == DiffCode.LEFTONLY:
                self.__main_window_ui.leftFileTextArea.insert('end', line['line'] + '\n', 'red')
                self.__main_window_ui.rightFileTextArea.insert('end', '\n', 'gray')
            elif line['code'] == DiffCode.CHANGED:
                for (i,c) in enumerate(line['line']):
                    self.__main_window_ui.leftFileTextArea.insert('end', c, 'darkred' if i in line['leftchanges'] else 'red')
                for (i,c) in enumerate(line['newline']):
                    self.__main_window_ui.rightFileTextArea.insert('end', c, 'darkgreen' if i in line['rightchanges'] else 'green')
                self.__main_window_ui.leftFileTextArea.insert('end', '\n')
                self.__main_window_ui.rightFileTextArea.insert('end', '\n')

            if line['code'] == DiffCode.LEFTONLY:
                self.__main_window_ui.leftLinenumbers.insert('end', str(leftlineno) + '\n', 'line')
                self.__main_window_ui.rightLinenumbers.insert('end', '\n', 'line')
                leftlineno += 1
            elif line['code'] == DiffCode.RIGHTONLY:
                self.__main_window_ui.leftLinenumbers.insert('end', '\n', 'line')
                self.__main_window_ui.rightLinenumbers.insert('end', str(rightlineno) + '\n', 'line')
                rightlineno += 1
            else:
                self.__main_window_ui.leftLinenumbers.insert('end', str(leftlineno) + '\n', 'line')
                self.__main_window_ui.rightLinenumbers.insert('end', str(rightlineno) + '\n', 'line')
                leftlineno += 1
                rightlineno += 1

        # calc width of line numbers texts and set it
        self.__main_window_ui.leftLinenumbers.config(width=len(str(leftlineno)))
        self.__main_window_ui.rightLinenumbers.config(width=len(str(rightlineno)))

        # disable text areas to prevent further editing
        self.__main_window_ui.leftFileTextArea.config(state=DISABLED)
        self.__main_window_ui.rightFileTextArea.config(state=DISABLED)
        self.__main_window_ui.leftLinenumbers.config(state=DISABLED)
        self.__main_window_ui.rightLinenumbers.config(state=DISABLED)

    def __cut(self):
        area = self.__getActiveTextArea()
        if area:
            area.event_generate("<<Cut>>")

    def __copy(self):
        area = self.__getActiveTextArea()
        if area:
            area.event_generate("<<Copy>>")

    def __paste(self):
        area = self.__getActiveTextArea()
        if area:
            area.event_generate("<<Paste>>")

    def __getActiveTextArea(self):
        if self.main_window.focus_get() == self.__main_window_ui.leftFileTextArea:
            return self.__main_window_ui.leftFileTextArea
        elif self.main_window.focus_get() == self.__main_window_ui.rightFileTextArea:
            return self.__main_window_ui.rightFileTextArea
        else:
            return None

    def __goToLine(self):
        line = askstring('Go to line', 'Enter line number:')
        if line:
            try:
                linenumber = int(line)
                self.__main_window_ui.leftFileTextArea.see(float(linenumber) + 5)
            except:
                pass

    def __startFindText(self):
        self.__main_window_ui.searchTextDialog.grid()
        self.__main_window_ui.searchTextDialog.focus()

    def __endFindText(self):
        self.__main_window_ui.leftFileTextArea.tag_remove('search', 1.0, END)
        self.__main_window_ui.rightFileTextArea.tag_remove('search', 1.0, END)
        self.__main_window_ui.searchTextDialog.unfocus()
        self.__main_window_ui.searchTextDialog.grid_remove()

    def __findNext(self, searchresult):
        term,leftpos,rightpos = searchresult['term'], searchresult['indices'][0], searchresult['indices'][1]
        if leftpos != -1:
            self.__main_window_ui.leftFileTextArea.tag_remove('search', 1.0, END)
            self.__main_window_ui.leftFileTextArea.tag_add('search', leftpos, '%s + %sc' % (leftpos, len(term)))
            # scroll to position plus five lines for visibility
            self.__main_window_ui.leftFileTextArea.see(float(leftpos) + 5)
        if rightpos != -1:
            self.__main_window_ui.rightFileTextArea.tag_remove('search', 1.0, END)
            self.__main_window_ui.rightFileTextArea.tag_add('search', rightpos, '%s + %sc' % (rightpos, len(term)))
            # scroll to position plus five lines for visibility
            self.__main_window_ui.rightFileTextArea.see(float(rightpos) + 5)

    def __exit(self):
        self.main_window.destroy()
